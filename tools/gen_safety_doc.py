"""Generate AgriRover_Safety_Mechanisms.docx in the repo root."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Heading colour helpers ────────────────────────────────────────────────────
def _set_heading_colour(style, rgb):
    style.font.color.rgb = RGBColor(*rgb)
    style.font.bold = True

_set_heading_colour(doc.styles['Heading 1'], (0x1a, 0x44, 0x7a))
_set_heading_colour(doc.styles['Heading 2'], (0x1a, 0x6a, 0x3a))
_set_heading_colour(doc.styles['Heading 3'], (0x44, 0x44, 0x44))

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(13)
doc.styles['Heading 3'].font.size = Pt(11)

# ── Utilities ─────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, hdr_hex='1A447A', alt_hex='EEF3FB'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # header
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c, hdr_hex)
    # data
    for ri, row_data in enumerate(rows):
        bg = alt_hex if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = tbl.rows[ri + 1].cells[ci]
            c.text = val
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(c, bg)
    doc.add_paragraph()

def hr():
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A447A')
    pb.append(bot)
    pPr.append(pb)

def centre(text, size=11, bold=False, rgb=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if rgb:
        r.font.color.rgb = RGBColor(*rgb)
    return p

def body(text, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
centre('AgriRover ROS2 System', size=24, bold=True, rgb=(0x1a, 0x44, 0x7a))
centre('Safety Mechanisms & Test Plan', size=18, bold=True, rgb=(0x1a, 0x6a, 0x3a))
doc.add_paragraph()
centre(f'Date: {datetime.date.today().strftime("%B %d, %Y")}', size=11, rgb=(0x55, 0x55, 0x55))
centre('Dual autonomous ground rover platform — precision agriculture', size=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1.  System Overview', level=1); hr()
body(
    'The AgriRover system consists of two autonomous ground rovers (RV1 master, '
    'RV2 slave), each driven by a Jetson Orin Nano running ROS2. Safety is implemented '
    'in four independent layers so that a failure in any single layer still brings the '
    'rovers to a safe neutral state.'
)
doc.add_paragraph()
body('Rover hardware:', italic=False)
bullet('RV1 (master) — Jetson Orin Nano, HM30 air unit, RP2040 with SBUS → SX1278 TX')
bullet('RV2 (slave)  — Jetson Orin Nano, RP2040 with SX1278 RX → PPM output')
bullet('GQC — SIYI MK32 Android ground control station, MAVLink over WiFi (sysid 255)')
doc.add_paragraph()
body('Safety layers (innermost to outermost):')
bullet('Layer 1 — RP2040 firmware  (hardware PPM at 50 Hz, fully independent of Jetson)')
bullet('Layer 2 — ROS2 nodes on Jetson  (navigator, rp2040_bridge, mavlink_bridge)')
bullet('Layer 3 — Android GQC application  (operator alerts and interlocks)')
bullet('Layer 4 — Physical RC hardware  (transmitter switch, RF link integrity)')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. MECHANISM INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2.  Safety Mechanism Inventory', level=1); hr()

# 2.1 ─────────────────────────────────────────────────────────────────────────
doc.add_heading('2.1  Layer 1 — RP2040 Firmware (Hardware PPM)', level=2)
body(
    'These mechanisms run inside the RP2040 microcontroller at 50 Hz, entirely '
    'independent of the Jetson and ROS2.  All result in PPM channels returning to '
    'neutral (1500 µs) unless otherwise stated.'
)
doc.add_paragraph()
add_table(
    ['#', 'Mechanism', 'Rover', 'Trigger', 'Effect'],
    [
        ['1', 'SBUS signal timeout',
         'Master (RV1)',
         'No SBUS frame for 200 ms\n(transmitter off or HM30 link lost)',
         'MODE_EMERGENCY → all PPM → 1500 µs\nLogs [SBUS_LOST]'],
        ['2', 'Emergency switch SWA (CH4)',
         'Both rovers',
         'sbus_ch[4] or rf_ch[4] < 1700 µs\n(physical 2-position switch flipped)',
         'MODE_EMERGENCY → all PPM → 1500 µs on both rovers simultaneously'],
        ['3', 'Jetson heartbeat timeout',
         'Both rovers',
         'No <HB:N> from rp2040_bridge for 300 ms',
         'hb_alive = false → AUTO blocked\nMODE_AUTO_NO_HB → PPM neutral'],
        ['4', 'Autonomous command timeout',
         'Both rovers',
         'No <J:c0..c7> from navigator for 500 ms',
         'cmd_fresh = false → MODE_AUTO_TIMEOUT → PPM neutral\n(rover coasts to stop)'],
        ['5', 'RF link timeout',
         'Slave (RV2)',
         'No LoRa packet for 500 ms',
         'MODE_EMERGENCY → all PPM → 1500 µs\nLogs [RF_LINK_LOST]'],
        ['6', 'SX1278 power-cycle watchdog',
         'Slave (RV2)',
         'RF lost + RegOpMode ≠ 0x85 every 2 s',
         'Automatic chip reinit + RX restart\nLogs [SX1278_REINIT opmode=0x...]'],
        ['7', 'PPM speed limit',
         'Both rovers',
         'Always active — all modes\n(MANUAL, AUTONOMOUS, RF relay)',
         'Throttle (CH1, PPM idx 0) and spin (CH2, PPM idx 1) hard-clamped to '
         '[1350–1650 µs] = ±30% of range.\nMK32 shows full range; hardware PPM capped.\n'
         'CH3–CH8 (switches, servos) unaffected.'],
    ]
)

# 2.2 ─────────────────────────────────────────────────────────────────────────
doc.add_heading('2.2  Layer 2 — ROS2 Nodes (Jetson)', level=2)
body(
    'These run on each Jetson Orin Nano. They depend on the RP2040 being responsive '
    'but are independent of GQC connectivity.'
)
doc.add_paragraph()
add_table(
    ['#', 'Mechanism', 'Node', 'Trigger', 'Effect'],
    [
        ['8', 'Heartbeat to RP2040',
         'rp2040_bridge',
         'Continuous 10 Hz loop',
         'Sends <HB:N> to keep hb_alive on RP2040.\nStopping this node triggers mechanism #3 within 300 ms.'],
        ['9', 'armed topic gate',
         'navigator',
         'armed topic = false',
         'Navigator stops publishing cmd_override.\nRP2040 reaches AUTO_TIMEOUT within 500 ms.'],
        ['10', 'mode topic gate',
         'navigator',
         'mode topic ≠ "AUTONOMOUS"',
         'Navigator does not compute or publish motion commands.'],
        ['11', 'RTK fix-loss safety',
         'navigator',
         'GPS fix type drops below RTK threshold',
         'Autonomous motion stopped.\nMission recording disabled.\nGQC RTK badge updated.'],
        ['12', 'Mission-complete auto-disarm',
         'mavlink_bridge',
         'wp_active = −1 published + no pending chunks',
         'Publishes MANUAL mode + DISARM.\nSTATUS badge drops to NA.\nFires only at true mission end.'],
    ]
)

# 2.3 ─────────────────────────────────────────────────────────────────────────
doc.add_heading('2.3  Layer 3 — Android GQC Application', level=2)
body(
    'Operator-facing interlocks and alerts. Critical commands use 3× UDP retry '
    '(100 ms spacing) to tolerate packet loss.'
)
doc.add_paragraph()
add_table(
    ['#', 'Mechanism', 'Trigger', 'Effect'],
    [
        ['13', 'E-STOP button',
         'Operator tap at any time',
         'DISARM (MAVCmd 400, p1=0) × 3 at 100 ms intervals to broadcast.\n'
         'Reaches both rovers simultaneously regardless of selection.'],
        ['14', 'DISARM before mission upload',
         'UPLOAD button pressed',
         'Sends DISARM × 3 before MISSION_COUNT.\n'
         'Prevents accidental motion during mission load.'],
        ['15', 'ARM safety gate (START button)',
         'START pressed in AUTO mode',
         'Blocks ARM + shows alert if:\n'
         '(a) No mission loaded — STATUS = NA\n'
         '(b) Rover STATUS = NA on rover side\n'
         '(c) Rover GPS > 0.5 m from WP[0]'],
        ['16', 'RC/HB link-mismatch auto-disarm',
         'CH9 switch ≠ slave HEARTBEAT for > 2 s',
         'checkLinkMismatch() detects inconsistency.\n'
         'Sends DISARM × 3 automatically. Logs mismatch event.'],
        ['17', 'HUD visual alerts',
         'Continuous telemetry monitoring',
         'SBUS dot (red=lost), RF dot (red=lost), RTK badge (colour=fix quality),\n'
         'STATUS badge (NA/MSL/ARM), HB dot blink.\n'
         'All initialised RED — require positive confirm to show green.'],
    ]
)

# 2.4 ─────────────────────────────────────────────────────────────────────────
doc.add_heading('2.4  Layer 4 — Physical RC Hardware', level=2)
doc.add_paragraph()
add_table(
    ['#', 'Mechanism', 'Trigger', 'Effect'],
    [
        ['18', 'Transmitter signal loss',
         'HM30 RF link drops or transmitter powered off',
         'SBUS signal lost on master RP2040 → mechanism #1 fires within 200 ms.'],
        ['19', 'SWA emergency switch (CH4)',
         'Operator flips 2-position switch on MK32',
         'Mechanism #2 fires simultaneously on both rovers via RP2040 firmware.\n'
         'Does not require Jetson or GQC connectivity.'],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TEST PLAN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3.  Safety Test Plan', level=1); hr()
body(
    'Each test is self-contained. "Pass condition" describes the observable outcome '
    'confirming correct behaviour. Tests are ordered from safest (bench / no motion) '
    'to full mission runs. Complete T1–T7 before the first outdoor drive.'
)
body(
    'Notation: "serial log" = USB CDC output from RP2040, readable with any serial '
    'terminal at 115200 baud or from monitor.py dashboard.',
    italic=True
)
doc.add_paragraph()

# helper for test tables
def test_table(rows):
    add_table(['Field', 'Details'], rows, hdr_hex='4A2E2E', alt_hex='FAF0F0')

doc.add_heading('T1 — SBUS Signal Timeout  [Mechanism #1]', level=2)
test_table([
    ['Prerequisites', 'Master RP2040 connected via USB. Rover on bench. Transmitter on, SBUS OK.'],
    ['Setup',         'Enable MANUAL mode. Apply slight throttle so PPM CH1 ≠ 1500.'],
    ['Action',        'Power off the MK32 transmitter (or disconnect HM30 cable).'],
    ['Observe',       'RP2040 serial log shows [SBUS_LOST] and MODE:EMERGENCY within 200 ms.'],
    ['Pass condition','CH1 and CH2 PPM return to 1500 µs within 200 ms. Mode = EMERGENCY.'],
    ['Restore',       'Power transmitter back on; SBUS_OK resumes automatically.'],
])

doc.add_heading('T2 — Emergency Switch SWA  [Mechanism #2]', level=2)
test_table([
    ['Prerequisites', 'Both rovers powered. Transmitter on with SBUS OK.'],
    ['Setup',         'MANUAL mode, slight motion on one rover.'],
    ['Action',        'Flip SWA (CH4) to the emergency/low position (< 1700 µs).'],
    ['Observe',       'Both rover serial logs show MODE:EMERGENCY. PPM → 1500 immediately.'],
    ['Pass condition','Both rovers stop while transmitter is still on and SBUS link is healthy.'],
    ['Restore',       'Flip SWA back to safe (high) position.'],
])

doc.add_heading('T3 — Jetson Heartbeat Timeout  [Mechanism #3]', level=2)
test_table([
    ['Prerequisites', 'Rover running full ROS2 stack. Armed in AUTO mode.'],
    ['Setup',         'Navigator running and sending motion commands.'],
    ['Action',        'Kill the rp2040_bridge node only (Ctrl-C or ros2 lifecycle shutdown).'],
    ['Observe',       'Within 300 ms: serial log shows AUTO-NO-HB; PPM returns to neutral.'],
    ['Pass condition','Rover stops within 300 ms. Navigator remains running and resumes when bridge restarts.'],
    ['Restore',       'Restart rp2040_bridge; heartbeat resumes; AUTO re-enables.'],
])

doc.add_heading('T4 — Autonomous Command Timeout  [Mechanism #4]', level=2)
test_table([
    ['Prerequisites', 'Rover armed and following a mission. rp2040_bridge running normally.'],
    ['Setup',         'Navigator actively publishing cmd_override at 10 Hz.'],
    ['Action',        'Kill the navigator node only (keep rp2040_bridge alive).'],
    ['Observe',       'Within 500 ms: serial log shows AUTO-TIMEOUT; PPM neutral.'],
    ['Pass condition','Rover stops. Bridge + heartbeat continue. RP2040 does NOT enter EMERGENCY.'],
    ['Restore',       'Restart navigator; re-arm via GQC START.'],
])

doc.add_heading('T5 — RF Link Timeout — Slave  [Mechanism #5]', level=2)
test_table([
    ['Prerequisites', 'Slave RP2040 running. RF link established. CH9 in slave (high) position.'],
    ['Setup',         'Slave moving slowly in RF mode.'],
    ['Action',        'Power off master RP2040 or block SX1278 antenna on master.'],
    ['Observe',       'Slave serial: [RF_LINK_LOST] then MODE:EMERGENCY within 500 ms.'],
    ['Pass condition','Slave stops within 500 ms. Master (RV1) is unaffected by this test.'],
    ['Restore',       'Power on master RP2040; RF link recovers automatically.'],
])

doc.add_heading('T6 — SX1278 Power-Cycle Watchdog  [Mechanism #6]', level=2)
test_table([
    ['Prerequisites', 'Slave RP2040 powered on bench. RF link established.'],
    ['Setup',         'Confirm [RF_LINK_OK] in serial log.'],
    ['Action',        'Momentarily cut power to SX1278 module only (keep RP2040 powered). Wait 3 s.'],
    ['Observe',       'Serial log: [RF_LINK_LOST] → [SX1278_REINIT opmode=0x…] → [RF_LINK_OK].'],
    ['Pass condition','RF link self-recovers without manual RP2040 reset. [SX1278_REINIT] line appears.'],
    ['Restore',       'No action needed; system self-heals.'],
])

doc.add_heading('T7 — PPM Speed Limit  [Mechanism #7]', level=2)
test_table([
    ['Prerequisites', 'RP2040 connected. Oscilloscope or ppm_decoder RP2040 reading GP15 PPM output.'],
    ['Setup',         'All modes can be tested; start with MANUAL.'],
    ['Action',        '(a) Push throttle stick to 100% (MK32 shows 2000 µs).\n'
                      '(b) Push spin stick to 100%.\n'
                      '(c) Move an auxiliary channel (e.g. CH5) to 2000 µs.'],
    ['Observe',       '(a)(b) Measured PPM pulse for CH1 and CH2 does not exceed 1650 µs.\n'
                      '(c) CH5 (PPM idx 4) reaches 2000 µs freely — unaffected.'],
    ['Pass condition','CH1 max = 1650 µs, CH2 max = 1650 µs. CH3–CH8 at full range. MK32 shows 2000 throughout.'],
    ['Restore',       'Return sticks to neutral.'],
])

doc.add_heading('T8 — armed Topic Gate  [Mechanism #9]', level=2)
test_table([
    ['Prerequisites', 'Rover running, armed in AUTO mode, navigator publishing commands.'],
    ['Setup',         'Rover in slow autonomous motion or bench-safe command stream.'],
    ['Action',        "Publish:  ros2 topic pub --once /rv1/armed std_msgs/Bool '{data: false}'"],
    ['Observe',       'Navigator stops publishing cmd_override. Within 500 ms: AUTO-TIMEOUT on RP2040. GQC STATUS: ARM → MSL.'],
    ['Pass condition','Rover stops. RP2040 shows AUTO-TIMEOUT (not EMERGENCY). Bridge still running.'],
    ['Restore',       'Re-arm via GQC START button.'],
])

doc.add_heading('T9 — RTK Fix-Loss Safety  [Mechanism #11]', level=2)
test_table([
    ['Prerequisites', 'Rover running outdoors with RTK corrections. GQC connected and showing RTK FIX (green badge).'],
    ['Setup',         'Rover armed in AUTO, RTK FIX confirmed.'],
    ['Action',        'Disconnect RTCM3 corrections feed (unplug rtk_forwarder) or block sky view to force fix drop.\n'
                      'In simulation: use simulator RTK quality controls.'],
    ['Observe',       'Navigator logs fix-loss warning. Rover stops autonomous motion. Recording disabled.\n'
                      'GQC RTK badge changes from RTK FIX to DGPS or 3D FIX.'],
    ['Pass condition','Rover stops autonomy. RTK badge in GQC turns non-green within a few seconds.'],
    ['Restore',       'Reconnect corrections; wait for RTK FIX before resuming.'],
])

doc.add_heading('T10 — Mission-Complete Auto-Disarm  [Mechanism #12]', level=2)
test_table([
    ['Prerequisites', 'Rover armed, short mission uploaded (3 waypoints minimum), GQC connected.'],
    ['Setup',         'START pressed; rover following mission at 30% speed limit.'],
    ['Action',        'Let the rover complete the full mission naturally — do not intervene.'],
    ['Observe',       'At final waypoint: wp_active = −1 published. GQC STATUS: ARM → NA. WP counter shows complete. Rover stops.'],
    ['Pass condition','Auto-disarm fires exactly once at true mission end (not at chunk boundaries). No manual action needed.'],
    ['Restore',       'Upload new mission to resume.'],
])

doc.add_heading('T11 — E-STOP Button  [Mechanism #13]', level=2)
test_table([
    ['Prerequisites', 'Rover armed and moving in AUTO mode. GQC connected via WiFi.'],
    ['Setup',         'Normal autonomous mission in progress.'],
    ['Action',        'Tap the E-STOP button in GQC.'],
    ['Observe',       'Three DISARM packets sent 100 ms apart (visible in monitor.py or Wireshark). Rover stops within 300 ms. STATUS → NA.'],
    ['Pass condition','Rover stops. Both rovers receive DISARM (broadcast address). 3 packets confirmed in network capture.'],
    ['Restore',       'Re-arm via START after re-uploading mission.'],
])

doc.add_heading('T12 — ARM Safety Gate  [Mechanism #15]', level=2)
body('Three sub-tests, all using the GQC START button:')
doc.add_paragraph()
add_table(
    ['Sub-test', 'Condition', 'Expected result'],
    [
        ['T12a', 'No mission loaded on rover (STATUS = NA)',
         'Toast message: "No mission on rover" — ARM blocked.'],
        ['T12b', 'Mission loaded but rover GPS > 0.5 m from WP[0]',
         'AlertDialog shows measured distance — ARM blocked.'],
        ['T12c', 'Mission loaded, rover within 0.5 m of WP[0]',
         'ARM proceeds normally; rover starts mission.'],
    ]
)

doc.add_heading('T13 — RC/HB Link-Mismatch Auto-Disarm  [Mechanism #16]', level=2)
test_table([
    ['Prerequisites', 'Both rovers connected. GQC running.'],
    ['Setup',         'Move CH9 switch to slave position (> 1750 µs). RV2 powered off or disconnected.'],
    ['Action',        'Wait > 2 s.'],
    ['Observe',       'GQC checkLinkMismatch() detects switch = slave but no RV2 HEARTBEAT. DISARM × 3 sent automatically.'],
    ['Pass condition','Auto-disarm fires within 2 s of mismatch detection. GQC logs the event.'],
    ['Restore',       'Return CH9 to neutral/master position. Power on RV2.'],
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. RECOMMENDED TEST ORDER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4.  Recommended Test Order for First Real-Rover Session', level=1); hr()
body(
    'Complete all bench tests (no rover motion) before any outdoor drive tests. '
    'Each test leaves the system safe for the next one.'
)
doc.add_paragraph()
add_table(
    ['Order', 'Test', 'Location', 'Rover motion'],
    [
        ['1',  'T7  — PPM speed limit',                'Bench (oscilloscope / ppm_decoder)', 'None'],
        ['2',  'T2  — SWA emergency switch',            'Bench or static field',              'None'],
        ['3',  'T1  — SBUS signal timeout',             'Bench or static field',              'None'],
        ['4',  'T6  — SX1278 power-cycle watchdog',     'Bench',                              'None'],
        ['5',  'T11 — E-STOP button',                   'Static field',                       'None'],
        ['6',  'T12 — ARM safety gate (a, b, c)',        'Static field',                       'None'],
        ['7',  'T5  — RF link timeout (slave)',          'Static field',                       'Slow crawl'],
        ['8',  'T8  — armed topic gate',                 'Open area',                          'Slow crawl'],
        ['9',  'T4  — navigator kill (cmd timeout)',     'Open area',                          'Slow crawl'],
        ['10', 'T3  — bridge kill (HB timeout)',         'Open area',                          'Slow crawl'],
        ['11', 'T13 — link-mismatch auto-disarm',        'Open area',                          'Optional'],
        ['12', 'T10 — mission-complete auto-disarm',     'Open area',                          'Full short mission'],
        ['13', 'T9  — RTK fix-loss safety',              'Outdoors with RTK corrections',      'Full mission'],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. QUICK REFERENCE — CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5.  Quick Reference — Safety Constants', level=1); hr()
add_table(
    ['Constant', 'Value', 'Where defined', 'Controls'],
    [
        ['SBUS_TIMEOUT_MS',        '200 ms',   'master/main.cpp',                    'Max gap before SBUS loss → EMERGENCY'],
        ['HB_TIMEOUT_MS',          '300 ms',   'master + slave main.cpp',            'Max gap before Jetson HB loss blocks AUTO'],
        ['CMD_TIMEOUT_MS',         '500 ms',   'master + slave main.cpp',            'Max gap before autonomous command timeout'],
        ['RF_TIMEOUT_MS',          '500 ms',   'slave/main.cpp',                     'Max gap before RF loss → EMERGENCY (slave)'],
        ['SX_WATCHDOG_MS',         '2000 ms',  'slave/main.cpp',                     'SX1278 health check interval when RF lost'],
        ['EMERGENCY_THRESH',       '1700 µs',  'master + slave main.cpp',            'SWA (CH4) threshold for emergency trigger'],
        ['SPEED_LIMIT_MAX',        '1650 µs',  'master + slave main.cpp',            'Max throttle/spin hardware PPM (30% limit)'],
        ['SPEED_LIMIT_MIN',        '1350 µs',  'master + slave main.cpp',            'Min throttle/spin hardware PPM (30% limit)'],
        ['ARM safety radius',      '0.5 m',    'Android GQC — RoverPositionManager', 'Max distance from WP[0] to allow ARM'],
        ['Mismatch timeout',       '2 s',      'Android GQC — MainActivity',         'CH9 vs HEARTBEAT mismatch before auto-disarm'],
        ['default_acceptance_radius', '0.1 m', 'rover1/2_params.yaml',              'Waypoint arrival radius in navigator'],
    ]
)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AgriRover ROS2 System — Safety Documentation')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
import os, sys
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'AgriRover_Safety_Mechanisms.docx')
out = os.path.normpath(out)
doc.save(out)
print(f'Saved: {out}')
